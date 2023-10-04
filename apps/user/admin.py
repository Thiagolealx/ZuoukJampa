from os import path

from django.contrib import admin
from django.shortcuts import render
from django.utils.html import format_html

from .forms import CongressistaFormAdmin
from .models import Lote, Categoria, Congressista, Pagamento, Entrada, Saida, Caixa
from django.utils.translation import gettext_lazy as _
from django.db.models import F
from django.db.models import Sum
from django.forms.models import BaseModelForm
from django import forms
from django.forms.models import BaseInlineFormSet
import docx
import openpyxl
from openpyxl.utils import get_column_letter
from django.http import HttpResponse

from decimal import Decimal

class LoteAdmin(admin.ModelAdmin):

    list_display = [
        "descricao",
        "valor_unitario",
    ]
    ordering = ["valor_unitario"]

admin.site.register(Lote, LoteAdmin)

class CategoriaAdmin(admin.ModelAdmin):

    search_fields = ['tipo']
    list_display = [
        "tipo",
        "sigla",
    ]
    ordering = ["tipo"]

admin.site.register(Categoria, CategoriaAdmin)

class LancamentoParcelaForm(forms.ModelForm, BaseModelForm):
    class Meta:
        model = Pagamento
        fields = ['congressista', 'valor_parcela','tipo',"numero_da_parcela"]

class PagamentoInline(admin.TabularInline):
    model = Pagamento
    form = LancamentoParcelaForm
    extra = 1
    max_num = 5
    can_delete = True


class PagamentoInlineFormSet(BaseInlineFormSet):
    def get_queryset(self):
        qs = super().get_queryset()
        return qs.order_by('-data_pagamento')

    def save_new(self, form, commit=True):
        obj = super().save_new(form, commit=False)
        obj.save()
        return obj
    


class CongressistaAdmin(admin.ModelAdmin):
    change_form_template = "congressita/change_form_congressista.html"
    list_display = [
        "nome_completo",
        "categoria",
        "ano",
        "uf",
        "lote",
        "get_valor_restante",
        "exibir_status_pagamento",
        "proxima_parcela",
        "numero_parcelas",
    ]
    list_filter = [
        "nome_completo",
        "cpf",        
        "uf",
        "categoria",          
     #"status_pagamento",
        ]
    list_select_related = ['lote']
    
    form = CongressistaFormAdmin
    change_list_template = "congressita/change_list_congressitas.html"
    inlines = [PagamentoInline]
    readonly_fields = ['valor_total_parcelas', 'get_valor_restante']
    list_max_show_all = 20


    def get_valor_restante(self, obj):
        return obj.get_valor_restante()
    get_valor_restante.short_description = 'Valor Restante'

    def exibir_status_pagamento(self, obj):
        valor_restante = obj.get_valor_restante()
        if abs(valor_restante) < 0.01:  # Ajuste a tolerância conforme necessário
            return format_html('<span style="color: green;">Pago</span>')
        else:
            return format_html('<span style="color: red;">Pendente</span>')

    exibir_status_pagamento.short_description = 'Status de Pagamento'



    def status_pagamento(self, obj):
        if obj.get_valor_restante() == 0:
            return 'Pago'
        else:
            return 'Pendente'

    status_pagamento.admin_order_field = 'valor_restante'
    status_pagamento.short_description = 'Status de Pagamento'    

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        queryset = queryset.annotate(total_lote=Sum('lote__valor_unitario'))
        return queryset

    def get_total_lote(self, obj):
        return Congressista.objects.all().aggregate(Sum('lote__valor_unitario'))['lote__valor_unitario__sum']


    get_total_lote.admin_order_field = 'total_lote'
    get_total_lote.short_description = 'Total Lote'

    tabs = ("ContratoPessoa",)
    save_on_top = True

    def get_total_parcelas(self, obj):
        total_parcelas = Congressista.objects.aggregate(total=Sum('pagamento__valor_parcela'))['total']
        return total_parcelas or 0

    def save(self, *args, **kwargs):
        self.total_parcelas = self.get_total_parcelas()
        super().save(*args, **kwargs)

    get_total_parcelas.admin_order_field = 'valor_total_parcelas'
    get_total_parcelas.short_description = 'Total Parcelas'

    def changelist_view(self, request, extra_context=None):
        response = super().changelist_view(request, extra_context=extra_context)
        try:
            qs = response.context_data["cl"].queryset
        except (AttributeError, KeyError):
            return response

        total_lote = Congressista.objects.aggregate(total_lote=Sum('lote__valor_unitario'))['total_lote'] or 0
        get_total_parcelas = Congressista.objects.aggregate(total_parcelas=Sum('pagamento__valor_parcela'))[
            'total_parcelas']

        response.context_data["total_lote"] = total_lote
        response.context_data["get_total_parcelas"] = get_total_parcelas

        return response
# Criação do relatorio

    actions = ['gerar_relatorio_word', 'gerar_relatorio_excel']
    def gerar_relatorio_word(self, request, queryset):
        # Obtém a UF selecionada do request GET
        uf_filter = request.GET.get('uf')

        # Crie um documento Word
        doc = docx.Document()

        # Adicione um título
        doc.add_heading('Relatório de Congressistas', level=1)

        # Filtra os congressistas com base na UF selecionada
        congressistas = queryset.filter(uf=uf_filter) if uf_filter else queryset

        # Ordene os congressistas por nome completo
        congressistas = congressistas.order_by('nome_completo')

        # Inicialize um contador
        contador = 1

        # Adicione os detalhes dos congressistas selecionados
        for congressista in congressistas:
            paragrafo = doc.add_paragraph()
            paragrafo.add_run(f'  {contador}')
            paragrafo.add_run(f' {congressista.nome_completo}')
            # paragrafo.add_run(f'\nUF: {congressista.uf}')
            contador += 1
            # paragrafo.add_run('\n' + '=' * 40)  # Linha separadora entre os registros

        # Crie uma resposta HTTP com o conteúdo do documento Word
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="relatorio_congressistas.docx"'
        doc.save(response)

        return response

    gerar_relatorio_word.short_description = "Gerar Relatório Word"

    def gerar_relatorio_excel(self, request, queryset):
        # Obtém a UF selecionada do request GET
        uf_filter = request.GET.get('uf')

        # Crie um novo workbook do Excel
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Relatório de Congressistas'

        # Crie os cabeçalhos das colunas
        headers = ['Número', 'Nome Completo', 'UF']
        for col_num, header in enumerate(headers, 1):
            col_letter = get_column_letter(col_num)
            ws[f"{col_letter}1"] = header

        # Filtra os congressistas com base na UF selecionada
        congressistas = queryset.filter(uf=uf_filter) if uf_filter else queryset

        # Ordene os congressistas por nome completo
        congressistas = congressistas.order_by('nome_completo')

        # Preencha os dados dos congressistas
        for row_num, congressista in enumerate(congressistas, 2):
            ws[f"A{row_num}"] = row_num - 1  # Número
            ws[f"B{row_num}"] = congressista.nome_completo  # Nome Completo
            ws[f"C{row_num}"] = congressista.uf  # UF
            ws[f"D{row_num}"] = congressista.lote.valor_unitario  # Valor

        # Crie uma resposta HTTP com o conteúdo do arquivo Excel
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="relatorio_congressistas.xlsx"'
        wb.save(response)

        return response

    gerar_relatorio_excel.short_description = "Gerar Relatório Excel"


class Media:
    js = ("admin/js/jquery.mask.min.js", "admin/js/custon.js", "jquery.js","admin/js/desativar_fka_pessoa.js")


admin.site.register(Congressista, CongressistaAdmin)

class PagamentoAdmin(admin.ModelAdmin):

    form = LancamentoParcelaForm
    search_fields = ['tipo']
    list_display = [
        "congressista",
        "data_pagamento",
        "valor_parcela",
        "numero_da_parcela",
        "get_valor_lote",
    ]
    ordering = ["congressista"]

    def get_valor_lote(self, obj):
        return obj.congressista.lote.valor_unitario

    get_valor_lote.short_description = 'Valor do Lote'

admin.site.register(Pagamento, PagamentoAdmin)

class StatusPagamentoFilter(admin.SimpleListFilter):
    title = _('Status de Pagamento')
    parameter_name = 'status_pagamento'

    def lookups(self, request, model_admin):
        return (
            ('pendente', _('Pendente')),
            ('pago', _('Pago')),
        )

    def queryset(self, request, queryset):
        if self.value() == 'pendente':
            return queryset.filter(valor_restante__gt=0)
        if self.value() == 'pago':
            return queryset.filter(valor_restante=0)


class EntradaAdmin(admin.ModelAdmin):
    list_display = ["descricao", "quantidade", "ano", "nome_empresa","valor_total_entrada"]
    list_filter = ["ano"]
    search_fields = ["descricao", "nome_empresa"]
    ordering = ["descricao"]
    change_list_template = "congressita/change_list_entradas.html"


    def save_model(self, request, obj, form, change):
        obj.valor_total_entrada = obj.valor_unitario * obj.quantidade
        obj.save()

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        queryset = queryset.annotate(total_saidas=Sum('valor_total_entrada'))
        return queryset

    def get_total_entradas(self, obj):
        return Entrada.objects.all().aggregate(Sum('valor_total_entrada'))['valor_total_entrada__sum']

    get_total_entradas.admin_order_field = 'get_total_entradas'
    get_total_entradas.short_description = 'Total Entradas'

    def changelist_view(self, request, extra_context=None):
        response = super().changelist_view(request, extra_context=extra_context)
        try:
            qs = response.context_data["cl"].queryset
        except (AttributeError, KeyError):
            return response
        get_total_entradas = (
            Entrada.objects.all().aggregate(Sum('valor_total_entrada'))['valor_total_entrada__sum']
        )

        response.context_data["get_total_entradas"] = get_total_entradas

        return response

    class Media:
        js = ("admin/js/jquery.mask.min.js", "admin/js/custon.js", "jquery.js", "admin/js/desativar_fka_pessoa.js")


class SaidaAdmin(admin.ModelAdmin):
    list_display = ["descricao", "quantidade", "ano", "nome_empresa","valor_total_saida" ]
    list_filter = ["ano"]
    search_fields = ["descricao", "nome_empresa"]
    ordering = ["descricao"]
    change_list_template = "congressita/change_list_saidas.html"

    def save_model(self, request, obj, form, change):
        obj.valor_total_saida = obj.valor_unitario * obj.quantidade
        obj.save()

    def get_queryset(self, request):
        queryset = super().get_queryset(request)
        queryset = queryset.annotate(total_saidas=Sum('valor_total_saida'))
        return queryset

    def get_total_saida(self, obj):
        return Saida.objects.all().aggregate(Sum('valor_total_saida'))['valor_total_saida__sum']

    get_total_saida.admin_order_field = 'get_total_saida'
    get_total_saida.short_description = 'Total Saidas'

    def changelist_view(self, request, extra_context=None):
        response = super().changelist_view(request, extra_context=extra_context)
        try:
            qs = response.context_data["cl"].queryset
        except (AttributeError, KeyError):
            return response
        get_total_saida = (
            Saida.objects.all().aggregate(Sum('valor_total_saida'))['valor_total_saida__sum']
        )

        response.context_data["get_total_saida"] = get_total_saida

        return response

    class Media:
        js = ("admin/js/jquery.mask.min.js", "admin/js/custon.js", "jquery.js", "admin/js/desativar_fka_pessoa.js")


admin.site.register(Entrada, EntradaAdmin)
admin.site.register(Saida, SaidaAdmin)




class CaixaAdmin(admin.ModelAdmin):
   
    readonly_fields = ['congressitas', 'entradas','saidas','saldo']
    change_list_template = "congressita/change_list_caixa.html"
    list_display = ['congressitas', 'entradas','saidas','saldo']
    
    def has_add_permission(self, request, obj=None):
        return False


    def changelist_view(self, request, extra_context=None):
            response = super().changelist_view(request, extra_context=extra_context)
            try:
                qs = response.context_data["cl"].queryset
            except (AttributeError, KeyError):
                return response

            total_lote = Congressista.objects.aggregate(total_lote=Sum('lote__valor_unitario'))['total_lote'] or 0
            get_total_parcelas = Congressista.objects.aggregate(total_parcelas=Sum('pagamento__valor_parcela'))[
                'total_parcelas']
            get_total_entradas= (
                Entrada.objects.all().aggregate(Sum('valor_total_entrada'))['valor_total_entrada__sum']
            )
            get_total_saida = (
            Saida.objects.all().aggregate(Sum('valor_total_saida'))['valor_total_saida__sum']
        )

            saldo = Decimal(get_total_parcelas or 0) + Decimal(get_total_entradas or 0) - Decimal(get_total_saida or 0)

            
        

            response.context_data["total_lote"] = total_lote
            response.context_data["get_total_parcelas"] = get_total_parcelas
            response.context_data["get_total_entradas"] = get_total_entradas
            response.context_data["get_total_saida"] = get_total_saida    
            response.context_data["saldo"] = saldo  
    

            return response





admin.site.register(Caixa, CaixaAdmin)